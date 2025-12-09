"""
Word Report Generator for PI Calibration Evidence Evaluation.

Generates professional Word documents (.docx) with:
- Title page with metadata
- Executive summary with pass/fail/needs-more counts
- Elements requiring attention section
- Per-element breakdown with evidence excerpts
"""

import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Union
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ReportGenerator:
    """Generate Word document evaluation report."""
    
    # Status icons (Unicode)
    STATUS_ICONS = {
        'Pass': '✓',
        'Fail': '✗',
        'Needs More Evidence': '?',
        'Error': '!'
    }
    
    # Status colors (RGB)
    STATUS_COLORS = {
        'Pass': RGBColor(0, 128, 0),       # Green
        'Fail': RGBColor(192, 0, 0),       # Red
        'Needs More Evidence': RGBColor(255, 165, 0),  # Orange
        'Error': RGBColor(128, 128, 128)   # Gray
    }
    
    def __init__(self, options: Optional[Dict] = None):
        """
        Initialize report generator with options.
        
        Args:
            options: Dict with keys:
                - include_pass: bool (default True)
                - include_fail: bool (default True)
                - include_needs_more: bool (default True)
                - include_excerpts: bool (default True)
                - include_reasoning: bool (default True)
                - include_description: bool (default True)
                - max_excerpt_length: int (default 500)
        """
        self.options = options or {}
        self.options.setdefault('include_pass', True)
        self.options.setdefault('include_fail', True)
        self.options.setdefault('include_needs_more', True)
        self.options.setdefault('include_excerpts', True)
        self.options.setdefault('include_reasoning', True)
        self.options.setdefault('include_description', True)
        self.options.setdefault('max_excerpt_length', 500)
        
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Configure document styles."""
        # Set default font
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Title style
        title_style = self.doc.styles['Title']
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading 1 style
        h1 = self.doc.styles['Heading 1']
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading 2 style
        h2 = self.doc.styles['Heading 2']
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading 3 style
        h3 = self.doc.styles['Heading 3']
        h3.font.size = Pt(12)
        h3.font.bold = True
    
    def _add_horizontal_line(self):
        """Add a horizontal line to the document."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '003366')
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    def generate(
        self,
        evaluation_results: Dict,
        matched_evidence: Optional[Dict] = None,
        evidence_files: Optional[List[str]] = None,
        output_path: Optional[Union[str, Path, BytesIO]] = None
    ) -> Union[str, BytesIO]:
        """
        Generate complete evaluation report.
        
        Args:
            evaluation_results: Dict with 'metadata', 'statistics', 'results' keys
            matched_evidence: Optional dict with matched evidence details
            evidence_files: List of evidence file names
            output_path: File path or BytesIO buffer to write to
            
        Returns:
            Output path string or BytesIO buffer
        """
        # Extract metadata
        metadata = evaluation_results.get('metadata', {})
        statistics = evaluation_results.get('statistics', {})
        results = evaluation_results.get('results', [])
        
        model = metadata.get('model', 'Unknown')
        timestamp = metadata.get('evaluation_timestamp', datetime.now().isoformat())
        
        # Build matched evidence lookup if provided
        evidence_lookup = {}
        if matched_evidence:
            for elem in matched_evidence.get('matched_elements', []):
                evidence_lookup[str(elem.get('PI-Element', ''))] = elem
        
        # Generate report sections
        self._add_title_page(model, evidence_files or [], timestamp)
        self._add_executive_summary(statistics, results)
        self._add_attention_section(results)
        self._add_detailed_breakdown(results, evidence_lookup)
        
        # Save document
        if output_path is None:
            output_path = BytesIO()
        
        self.doc.save(output_path)
        
        if isinstance(output_path, BytesIO):
            output_path.seek(0)
        
        return output_path
    
    def _add_title_page(self, model: str, evidence_files: List[str], timestamp: str):
        """Add report title and metadata."""
        # Title
        title = self.doc.add_heading('PI Calibration Evidence Evaluation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Metadata table
        table = self.doc.add_table(rows=3, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Parse timestamp
        try:
            dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
            formatted_date = dt.strftime('%B %d, %Y %I:%M %p')
        except:
            formatted_date = timestamp
        
        cells = [
            ('Generated:', formatted_date),
            ('Model:', model),
            ('Evidence Files:', ', '.join(evidence_files) if evidence_files else 'N/A'),
        ]
        
        for i, (label, value) in enumerate(cells):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value
        
        self.doc.add_paragraph()
        self._add_horizontal_line()
        self.doc.add_page_break()
    
    def _add_executive_summary(self, statistics: Dict, results: List[Dict]):
        """Add summary statistics section."""
        self.doc.add_heading('Executive Summary', level=1)
        
        total = statistics.get('total', len(results))
        pass_count = statistics.get('pass', 0)
        fail_count = statistics.get('fail', 0)
        needs_count = statistics.get('needs_more_evidence', 0)
        error_count = statistics.get('error', 0)
        
        # Summary paragraph
        p = self.doc.add_paragraph()
        p.add_run(f'Total Elements Evaluated: ').bold = True
        p.add_run(f'{total}')
        
        self.doc.add_paragraph()
        
        # Statistics table
        table = self.doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        headers = ['Status', 'Count', 'Percentage']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            self._shade_cell(cell, RGBColor(0, 51, 102))
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        rows_data = [
            ('✓ Pass', pass_count, pass_count / total * 100 if total > 0 else 0, 'Pass'),
            ('? Needs More Evidence', needs_count, needs_count / total * 100 if total > 0 else 0, 'Needs More Evidence'),
            ('✗ Fail', fail_count, fail_count / total * 100 if total > 0 else 0, 'Fail'),
            ('! Error', error_count, error_count / total * 100 if total > 0 else 0, 'Error'),
        ]
        
        for i, (status, count, pct, status_key) in enumerate(rows_data, start=1):
            row = table.rows[i]
            row.cells[0].text = status
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = self.STATUS_COLORS.get(status_key, RGBColor(0, 0, 0))
            row.cells[1].text = str(count)
            row.cells[2].text = f'{pct:.1f}%'
        
        self.doc.add_paragraph()
        self._add_horizontal_line()
    
    def _shade_cell(self, cell, color: RGBColor):
        """Apply background shading to a table cell."""
        shading = OxmlElement('w:shd')
        # RGBColor is indexable: color[0]=red, color[1]=green, color[2]=blue
        shading.set(qn('w:fill'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _add_attention_section(self, results: List[Dict]):
        """Add section highlighting failed and needs-more elements."""
        self.doc.add_heading('Elements Requiring Attention', level=1)
        
        # Failed elements
        failed = [r for r in results if r.get('Status', '').lower() == 'fail']
        if failed:
            p = self.doc.add_paragraph()
            run = p.add_run('✗ FAIL')
            run.bold = True
            run.font.color.rgb = self.STATUS_COLORS['Fail']
            p.add_run(f' ({len(failed)} elements)')
            
            for r in failed:
                elem_id = r.get('PI-Element', 'Unknown')
                self.doc.add_paragraph(f'• Element {elem_id}', style='List Bullet')
        else:
            p = self.doc.add_paragraph()
            run = p.add_run('✓ No failed elements')
            run.font.color.rgb = self.STATUS_COLORS['Pass']
        
        self.doc.add_paragraph()
        
        # Needs more evidence elements
        needs_more = [r for r in results if 'needs' in r.get('Status', '').lower()]
        if needs_more:
            p = self.doc.add_paragraph()
            run = p.add_run('? NEEDS MORE EVIDENCE')
            run.bold = True
            run.font.color.rgb = self.STATUS_COLORS['Needs More Evidence']
            p.add_run(f' ({len(needs_more)} elements)')
            
            for r in needs_more:
                elem_id = r.get('PI-Element', 'Unknown')
                self.doc.add_paragraph(f'• Element {elem_id}', style='List Bullet')
        else:
            p = self.doc.add_paragraph()
            run = p.add_run('✓ All elements have sufficient evidence')
            run.font.color.rgb = self.STATUS_COLORS['Pass']
        
        self.doc.add_paragraph()
        self._add_horizontal_line()
        self.doc.add_page_break()
    
    def _add_detailed_breakdown(self, results: List[Dict], evidence_lookup: Dict):
        """Add per-element detailed breakdown."""
        self.doc.add_heading('Detailed Element Breakdown', level=1)
        
        for result in results:
            status = result.get('Status', 'Unknown')
            
            # Skip based on options
            if status.lower() == 'pass' and not self.options['include_pass']:
                continue
            if status.lower() == 'fail' and not self.options['include_fail']:
                continue
            if 'needs' in status.lower() and not self.options['include_needs_more']:
                continue
            
            self._add_element_section(result, evidence_lookup)
    
    def _add_element_section(self, result: Dict, evidence_lookup: Dict):
        """Add single element breakdown."""
        elem_id = str(result.get('PI-Element', 'Unknown'))
        status = result.get('Status', 'Unknown')
        reasoning = result.get('LLM response', '')
        slide_nums = result.get('Evidence Slide num', [])
        
        # Element header
        self.doc.add_heading(f'Element {elem_id}', level=2)
        
        # Status with color
        p = self.doc.add_paragraph()
        p.add_run('Status: ').bold = True
        
        icon = self.STATUS_ICONS.get(status, '?')
        status_run = p.add_run(f'{icon} {status.upper()}')
        status_run.bold = True
        
        # Determine status key for color
        if status.lower() == 'pass':
            status_key = 'Pass'
        elif status.lower() == 'fail':
            status_key = 'Fail'
        elif 'needs' in status.lower():
            status_key = 'Needs More Evidence'
        else:
            status_key = 'Error'
        
        status_run.font.color.rgb = self.STATUS_COLORS.get(status_key, RGBColor(0, 0, 0))
        
        # Matching slides
        if slide_nums:
            p = self.doc.add_paragraph()
            p.add_run('Matching Slides: ').bold = True
            p.add_run(', '.join(str(s) for s in slide_nums))
        
        # Element description from evidence lookup
        if self.options['include_description'] and elem_id in evidence_lookup:
            elem_data = evidence_lookup[elem_id]
            instructions = elem_data.get('Calibrator instructions', {})
            ask_for = instructions.get('Ask/Look For', '')
            
            if ask_for:
                self.doc.add_heading('Element Description:', level=3)
                p = self.doc.add_paragraph()
                p.add_run(ask_for).italic = True
        
        # Evaluation reasoning
        if self.options['include_reasoning'] and reasoning:
            self.doc.add_heading('Evaluation Reasoning:', level=3)
            self.doc.add_paragraph(reasoning)
        
        # Evidence excerpts
        if self.options['include_excerpts'] and elem_id in evidence_lookup:
            elem_data = evidence_lookup[elem_id]
            evidence_list = elem_data.get('Evidence', [])
            
            if evidence_list:
                self.doc.add_heading('Evidence Text Excerpts:', level=3)
                
                for evidence in evidence_list[:3]:  # Limit to first 3 excerpts
                    slide_idx = evidence.get('slide_index', '?')
                    full_text = evidence.get('full_text', evidence.get('text_preview', ''))
                    
                    # Truncate if too long
                    max_len = self.options['max_excerpt_length']
                    if len(full_text) > max_len:
                        full_text = full_text[:max_len] + '...'
                    
                    # Add excerpt in bordered box style
                    p = self.doc.add_paragraph()
                    run = p.add_run(f'From Slide {slide_idx}:')
                    run.bold = True
                    run.font.size = Pt(10)
                    
                    # Add the excerpt text
                    excerpt_p = self.doc.add_paragraph()
                    excerpt_p.paragraph_format.left_indent = Inches(0.25)
                    excerpt_run = excerpt_p.add_run(full_text)
                    excerpt_run.font.size = Pt(10)
                    excerpt_run.font.color.rgb = RGBColor(64, 64, 64)
        
        # Add separator
        self._add_horizontal_line()


def generate_word_report(
    evaluation_results_path: Union[str, Path],
    output_path: Union[str, Path],
    matched_evidence_path: Optional[Union[str, Path]] = None,
    evidence_files: Optional[List[str]] = None,
    options: Optional[Dict] = None
) -> str:
    """
    Convenience function to generate Word report from JSON files.
    
    Args:
        evaluation_results_path: Path to evaluation_results.json
        output_path: Path for output .docx file
        matched_evidence_path: Optional path to matched_evidence.json
        evidence_files: List of evidence file names for metadata
        options: Report generation options
        
    Returns:
        Output file path
    """
    # Load evaluation results
    with open(evaluation_results_path, 'r', encoding='utf-8') as f:
        evaluation_results = json.load(f)
    
    # Load matched evidence if provided
    matched_evidence = None
    if matched_evidence_path:
        with open(matched_evidence_path, 'r', encoding='utf-8') as f:
            matched_evidence = json.load(f)
    
    # Generate report
    generator = ReportGenerator(options)
    generator.generate(
        evaluation_results=evaluation_results,
        matched_evidence=matched_evidence,
        evidence_files=evidence_files,
        output_path=output_path
    )
    
    return str(output_path)


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Generate Word report from evaluation results')
    parser.add_argument('--evaluation', '-e', required=True, help='Path to evaluation_results.json')
    parser.add_argument('--matched', '-m', help='Path to matched_evidence.json (optional)')
    parser.add_argument('--output', '-o', required=True, help='Output .docx file path')
    parser.add_argument('--evidence-files', '-f', nargs='+', help='Evidence file names')
    parser.add_argument('--no-pass', action='store_true', help='Exclude passed elements')
    parser.add_argument('--no-excerpts', action='store_true', help='Exclude evidence excerpts')
    
    args = parser.parse_args()
    
    options = {
        'include_pass': not args.no_pass,
        'include_excerpts': not args.no_excerpts,
    }
    
    output = generate_word_report(
        evaluation_results_path=args.evaluation,
        output_path=args.output,
        matched_evidence_path=args.matched,
        evidence_files=args.evidence_files,
        options=options
    )
    
    print(f'Report generated: {output}')
